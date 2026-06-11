import os
import random
import requests

from bs4 import BeautifulSoup
from dotenv import load_dotenv

from langchain_groq import ChatGroq
from langchain_ollama import OllamaEmbeddings

from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser

from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

# ==========================================
# LOAD ENV
# ==========================================
load_dotenv()

GROQ_API_KEY = os.getenv("GROQ_API_KEY")

# ==========================================
# EMBEDDINGS
# ==========================================
embedding_model = OllamaEmbeddings(
    model="nomic-embed-text"
)

# ==========================================
# TEXT SPLITTER
# ==========================================
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=500,
    chunk_overlap=50
)

# ==========================================
# LLM
# ==========================================
llm = ChatGroq(
    groq_api_key=GROQ_API_KEY,
    model_name="llama-3.3-70b-versatile",
    temperature=0.7
)

# ==========================================
# PROMPT
# ==========================================
prompt = ChatPromptTemplate.from_template("""
You are an AI Assistant.

Context:
{context}

Question:
{question}

Answer:
""")

rag_chain = (
    prompt
    | llm
    | StrOutputParser()
)

# ==========================================
# CREATE VECTOR DB
# ==========================================
def create_vector_db(texts):

    docs = [
        Document(page_content=t)
        for t in texts
    ]

    split_docs = text_splitter.split_documents(
        docs
    )

    vectorstore = FAISS.from_documents(
        split_docs,
        embedding_model
    )

    vectorstore.save_local("faiss_db")

# ==========================================
# LOAD VECTOR DB
# ==========================================
def load_vector_db():

    return FAISS.load_local(
        "faiss_db",
        embedding_model,
        allow_dangerous_deserialization=True
    )

# ==========================================
# RETRIEVE CONTEXT
# ==========================================
def retrieve_context(query):

    vectorstore = load_vector_db()

    docs = vectorstore.similarity_search(
        query,
        k=3
    )

    context = "\n".join([
        doc.page_content
        for doc in docs
    ])

    return context

# ==========================================
# READ TXT
# ==========================================
def read_txt(file_path):

    with open(
        file_path,
        "r",
        encoding="utf-8"
    ) as file:

        return file.read()

# ==========================================
# READ PDF
# ==========================================
def read_pdf(file_path):

    from PyPDF2 import PdfReader

    reader = PdfReader(file_path)

    text = ""

    for page in reader.pages:

        extracted = page.extract_text()

        if extracted:
            text += extracted

    return text

# ==========================================
# READ DOCX
# ==========================================
def read_docx(file_path):

    import docx

    doc = docx.Document(file_path)

    return "\n".join([
        para.text
        for para in doc.paragraphs
    ])

# ==========================================
# WEBSITE PROCESSING
# ==========================================
def process_website(url):

    response = requests.get(url)

    soup = BeautifulSoup(
        response.text,
        "html.parser"
    )

    return soup.get_text()

# ==========================================
# OBJECTIVE QUESTIONS
# ==========================================
def generate_objective_questions(
    topic,
    num_questions=10
):

    question_templates = [
        f"What is {topic}?",
        f"Which statement is correct about {topic}?",
        f"What is the purpose of {topic}?",
        f"Which feature belongs to {topic}?",
        f"Why is {topic} important?",
        f"Which option best describes {topic}?",
        f"What are the advantages of {topic}?",
        f"How is {topic} used?",
        f"What is a key concept in {topic}?",
        f"Which of the following relates to {topic}?"
    ]

    option_sets = [
        [
            f"{topic} Concept",
            "Database",
            "Programming Language",
            "Browser"
        ],
        [
            "Operating System",
            f"{topic} Technology",
            "Compiler",
            "Hardware Device"
        ],
        [
            "Networking Tool",
            "Cloud Service",
            f"{topic} Method",
            "Text Editor"
        ]
    ]

    questions = []

    for i in range(num_questions):

        q_text = random.choice(
            question_templates
        )

        options = random.choice(
            option_sets
        )

        correct_answer = options[0]

        random.shuffle(options)

        q = {
            "question": f"{q_text} ({i+1})",
            "options": options,
            "answer": correct_answer
        }

        questions.append(q)

    return questions

# ==========================================
# SUBJECTIVE QUESTIONS
# ==========================================
def generate_subjective_questions(
    topic,
    num_questions=10
):

    questions = []

    templates = [
        f"Explain {topic} in detail.",
        f"Discuss the advantages of {topic}.",
        f"Describe the applications of {topic}.",
        f"What are the challenges in {topic}?",
        f"Explain the importance of {topic}.",
    ]

    for i in range(num_questions):

        q = random.choice(templates)

        questions.append(
            f"{q} ({i+1})"
        )

    return questions

# ==========================================
# OBJECTIVE EVALUATION
# ==========================================
def evaluate_objective(
    user_answers,
    questions
):

    score = 0

    for user, q in zip(
        user_answers,
        questions
    ):

        if user == q["answer"]:

            score += 1

    return score

# ==========================================
# SUBJECTIVE EVALUATION
# ==========================================
def evaluate_subjective(
    question,
    answer
):

    evaluation_prompt = f"""
    Evaluate the following answer.

    Question:
    {question}

    Answer:
    {answer}

    Give:
    1. Score out of 10
    2. Strengths
    3. Improvements
    """

    response = llm.invoke(
        evaluation_prompt
    )

    return response.content

# ==========================================
# FEEDBACK
# ==========================================
def generate_feedback(score):

    if score >= 8:

        return {
            "performance": "Excellent",
            "feedback": [
                "Strong conceptual understanding",
                "Very good technical clarity"
            ]
        }

    elif score >= 5:

        return {
            "performance": "Good",
            "feedback": [
                "Improve technical depth",
                "Add more examples"
            ]
        }

    else:

        return {
            "performance": "Needs Improvement",
            "feedback": [
                "Study concepts carefully",
                "Practice more questions"
            ]
        }

# ==========================================
# AI CHATBOT
# ==========================================
def ask_ai(question):

    try:

        context = retrieve_context(question)

        response = rag_chain.invoke({
            "context": context,
            "question": question
        })

        return response

    except Exception as e:

        return str(e)